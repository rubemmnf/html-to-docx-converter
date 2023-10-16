from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches
import requests
import os
import re

def get_web_content(url):
    """Fetch the content of the web page."""
    response = requests.get(url)
    response.raise_for_status()
    return response.text

def download_image(img_url, download_folder="downloaded_images"):
    """Download an image and return its local path."""
    if not os.path.exists(download_folder):
        os.makedirs(download_folder)

    response = requests.get(img_url, stream=True)
    response.raise_for_status()

    local_filename = os.path.join(download_folder, os.path.basename(img_url))
    with open(local_filename, 'wb') as f:
        for chunk in response.iter_content(chunk_size=8192):
            f.write(chunk)

    return local_filename

def extract_images_and_process_content(html_content):
    """Extract images, download them, and process the content."""
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Placeholder for downloaded images
    images = []
    
    # Download images and replace img tags with placeholders in the content
    for img_tag in soup.find_all('img'):
        img_url = img_tag['src']
        local_path = download_image(img_url)
        images.append(local_path)
        img_tag.replace_with(f"[[IMAGE:{local_path}]]")
    
    # Process text content
    text = soup.get_text()
    
    text = text.replace(")\n", ") ")
    text = text.replace("\n(", "(")
 
    text = re.sub(r'\n\s*\n', '\n', text)
    
    return text, images

def save_to_docx(text, images, filename):
    doc = Document()
    section = doc.sections[0]
    sect_pr = section._sectPr
    cols = sect_pr.xpath('./w:cols')[0]
    cols.set(qn('w:num'),'2')
    column_width = Inches(3) 
    
    # Iterating through the processed text and inserting images at their respective placeholders
    for line in text.splitlines():
        if "[[IMAGE:" in line:
            img_path = line.split('[[IMAGE:')[1].split(']]')[0]
            doc.add_picture(img_path, width=column_width)
        else:
            line = line.replace("Questão ", "\nQuestão ")
            paragraph = doc.add_paragraph(line)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph.paragraph_format.space_after = Inches(0)  # Remove spacing after paragraph

    doc.save(filename)

def main():
    url = input("Please enter the URL of the web page: ")
    html_content = get_web_content(url)
    
    processed_text, extracted_images = extract_images_and_process_content(html_content)
    save_to_docx(processed_text, extracted_images, "output.docx")
    print("Content saved to output.docx")

if __name__ == "__main__":
    main()
